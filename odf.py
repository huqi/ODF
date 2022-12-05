import ctypes
import json
import os
import sys
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QGuiApplication, QIcon, QPixmap
from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox, QWidget
from win32com import client as wc

from odf_conf import *
from odf_ui import Ui_Form

ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("myappid")


def get_file_extension(file):
    return os.path.splitext(file)[-1]


def is_word_file(file):
    return any(file.endswith(ext) for ext in ['.doc', '.docx', 'DOC', '.DOCX'])


def doc_to_docx(file):
    word = wc.Dispatch("Word.Application")
    filename = "{}x".format(new_file(file))
    doc = word.Documents.Open(file)
    doc.SaveAs(filename, 12)
    doc.Close()  # 关闭原来word文件
    word.Quit()
    return filename


def get_time_str():
    return time.strftime("%Y%m%d%H%M%S", time.localtime())


def new_file(oldfile):
    path, file = os.path.split(oldfile)
    filename, ext = os.path.splitext(file)
    new_filename = get_time_str() + "-" + filename + ext

    new_path = os.path.join(path, new_filename)
    return '\\'.join(new_path.split('/'))


def get_conf(self):

    data = None

    try:
        data = load_conf(CONF_FILENAME)
    except:
        QMessageBox.warning(self, STR_WARNING, STR_LOAD_CONF_ERR)
        data = restore_conf(CONF_FILENAME, P_STYLE)

    return data


class Para:
    text = ""
    content = ""
    type = ""

    def __init__(self, text, type=None):
        if not type:
            self.type = self.get_type(text)
        else:
            self.type = type

        if self.type in ('head2', 'head3', 'head4'):
            self.text, self.content = text.split("。", 1)
            self.text += "。"
        else:
            self.text = text

    def get_type(self, text):
        for k, v in HEAD_SIGN.items():
            for sign in v:
                if text.startswith(sign):
                    return k

        return 'text'

    def print_type(self):
        print(self.type)

    def print_text(self):
        print(self.text)


class Doc:
    document = None
    filename = ""

    def __init__(self, file):
        self.doc = []
        self.paragraphs = []
        self.document = Document(file)
        self.filename = file
        self.step = 0
        self.step_max = 0
        self.d_to_dx_flag = False

    def print(self):
        for p in self.paragraphs:
            p.print_type()
            p.print_text()

    def setConf(self, b_title, a_title, tail):
        self.b_title_count = b_title
        self.a_title_count = a_title
        self.tail_count = tail

    def setAbility(self, ability):
        self.ability = ability

    def setSignal(self, sig_label, sig_label_info, sig_progress_init, sig_progress_index, sig_done):
        self.sig_label = sig_label
        self.sig_label_info = sig_label_info
        self.sig_progress_init = sig_progress_init
        self.sig_progress_index = sig_progress_index
        self.sig_done = sig_done

    def setDToDx(self, flag):
        self.d_to_dx_flag = flag

    def setStep(self, step, step_max):
        self.step = step
        self.step_max = step_max

    def wash_line(self):
        para = []
        self.step += 1
        self.sig_label.emit(self.step, self.step_max)
        self.sig_label_info.emit(STR_LABEL_REPLACE_BLANK)
        self.sig_progress_init.emit(len(self.doc))
        for i, line in enumerate(self.doc):
            self.sig_progress_index.emit(i+1)
            if len(line) > 0:
                para.append(line)

        return para

    def load_para(self, para):
        b_title_c = 0
        title_done = False
        a_title_c = 0
        target_line_done = False
        tail_line_c = 0
        tail_blank_line_done = False
        line_no = 0

        self.step += 1
        self.sig_label.emit(self.step, self.step_max)
        self.sig_label_info.emit(STR_LABEL_FORMAT)
        self.sig_progress_init.emit(len(para))

        for i, line in enumerate(para):
            line_no += 1
            self.sig_progress_index.emit(line_no)

            if self.ability['checkBox_b_title'] and b_title_c < self.b_title_count:
                self.paragraphs.append(Para(line, "b_title"))
                b_title_c += 1

                if b_title_c == self.b_title_count:
                    self.paragraphs.append(Para("", "text"))

                continue

            if not title_done:
                self.paragraphs.append(Para(line, "title"))
                if not self.ability['checkBox_a_title']:
                    self.paragraphs.append(Para("", "text"))

                title_done = True
                continue

            if self.ability['checkBox_a_title'] and a_title_c < self.a_title_count:
                self.paragraphs.append(Para(line, "a_title"))
                a_title_c += 1

                if a_title_c == self.a_title_count:
                    self.paragraphs.append(Para("", "text"))

                continue

            if self.ability['checkBox_target'] and not target_line_done:
                self.paragraphs.append(Para(line, "target"))
                target_line_done = True
                continue

            if self.ability['checkBox_tail'] and line_no > len(para) - self.tail_count and tail_line_c < self.tail_count:
                if not tail_blank_line_done:
                    self.paragraphs.append(Para("", "text"))
                    self.paragraphs.append(Para("", "text"))
                    tail_blank_line_done = True

                self.paragraphs.append(Para(line, "tail"))
                tail_line_c += 1
                continue

            self.paragraphs.append(Para(line))

    def replace_x(self, text):
        t = text
        for x in x_list:
            t = t.replace(x, '')

        return t

    def replace_blank(self):
        self.step += 1
        self.sig_label.emit(self.step, self.step_max)
        self.sig_label_info.emit(STR_LABEL_REPLACE_SPACE)
        self.sig_progress_init.emit(len(self.document.paragraphs))

        for i, line in enumerate(self.document.paragraphs):
            self.sig_progress_index.emit(i+1)
            self.doc.append(self.replace_x(line.text))

    def run(self):
        self.replace_blank()
        self.load_para(self.wash_line())

    def getTimeStr(self):
        return time.strftime("%Y%m%d%H%M%S", time.localtime())

    def write(self):
        doc = Document()

        self.step += 1
        self.sig_label.emit(self.step, self.step_max)
        self.sig_label_info.emit(STR_LABEL_WRITE)
        self.sig_progress_init.emit(len(self.paragraphs))

        for i, line in enumerate(self.paragraphs):
            self.sig_progress_index.emit(i+1)

            p = doc.add_paragraph()
            run = p.add_run(line.text)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(
                qn('w:eastAsia'), P_STYLE[line.type]['font'])
            run.font.size = Pt(FONT_SIZE_PT[P_STYLE[line.type]['font_size']])

            if line.content:
                run_content = p.add_run(line.content)
                run_content.font.name = 'Times New Roman'
                run_content.element.rPr.rFonts.set(
                    qn('w:eastAsia'), P_STYLE['text']['font'])
                run_content.font.size = Pt(
                    FONT_SIZE_PT[P_STYLE['text']['font_size']])

            p.paragraph_format.space_before = Pt(
                P_STYLE[line.type]['space_before'])
            p.paragraph_format.space_after = Pt(
                P_STYLE[line.type]['space_after'])
            p.paragraph_format.line_spacing = Pt(
                P_STYLE[line.type]['line_spacing'])
            p.paragraph_format.alignment = PARA_ALIGN[P_STYLE[line.type]['align']]
            p.paragraph_format.first_line_indent = run.font.size * \
                P_STYLE[line.type]['indent']

        if self.d_to_dx_flag:
            os.remove(self.filename)

        doc.save(self.filename)
        self.sig_done.emit(self.filename.split('\\')[-1])


class MyMainForm(QWidget, Ui_Form):
    def __init__(self, parent=None):
        '''
        构造函数
        '''
        super(MyMainForm, self).__init__(parent)
        self.setupUi(self)

        self.img_blue = QPixmap(IMG_DRAG_DROP_BLUE_PATH)
        self.img_grey = QPixmap(IMG_DRAG_DROP_GRAY_PATH)

        self.setWindowTitle(STR_WINDOW_TITLE)
        self.setWindowIcon(QIcon(IMG_ICON_PATH))

        x, y = self.getCenterPos()

        self.setGeometry(
            x, y, MAIN_FRAME_SIZE['WIDTH'], MAIN_FRAME_SIZE['HEIGHT'])
        self.setFixedSize(self.width(), self.height())
        self.label_button.setAlignment(Qt.AlignCenter)
        self.setAcceptDrops(True)
        self.eventRestore(None, None, None)

    def setConfig(self, conf):
        self.conf = conf
        self.checkBox_b_title.setChecked(
            self.conf['default']['b_title_checked'])
        self.checkBox_a_title.setChecked(
            self.conf['default']['a_title_checked'])
        self.checkBox_target.setChecked(self.conf['default']['target_checked'])
        self.checkBox_tail.setChecked(self.conf['default']['tail_checked'])

    def getCenterPos(self):
        '''
        获得居中坐标
        '''
        screen = QGuiApplication.primaryScreen().size()

        x = int((screen.width() - MAIN_FRAME_SIZE['WIDTH']) / 2)
        y = int((screen.height() - MAIN_FRAME_SIZE['HEIGHT']) / 2)

        return x, y

    def dragEnterEvent(self, event):
        '''
        拖入
        '''
        event.accept()
        self.label_button.setPixmap(self.img_blue)

    def dragMoveEvent(self, event):
        '''
        移动
        '''
        pass

    def dragLeaveEvent(self, event):
        '''
        移出
        '''
        self.label_button.setPixmap(self.img_grey)

    def showMessageBox(self, type, title, msg):
        if type == "warning":
            QMessageBox.warning(self, title, msg)
        elif type == "info":
            QMessageBox.information(self, title, msg)

    def eventRestore(self, type, title, message):
        '''
        进度条重置
        '''
        self.showMessageBox(type, title, message)
        self.label_button.setPixmap(self.img_grey)
        self.progressBar.setValue(0)
        self.label_progress.setText("(0/0)")
        self.label_info.setText("选择文档特征后拖入Word文件")
        self.running = False

    def evenErrorCallback(self):
        self.eventRestore("warning", STR_ERROR, STR_UNKNOWN)

    def evenLabelCallback(self, i, count):
        self.label_progress.setText("({0}/{1})".format(i, count))

    def evenLabelInfoCallback(self, s):
        self.label_info.setText(s)

    def evenProgressInitCallback(self, max):
        self.progressBar.setMaximum(max)

    def evenProgressIndexCallback(self, i):
        self.progressBar.setValue(i)

    def evenDoneCallback(self, filename):
        self.eventRestore("info", STR_DONE, STR_ALREADY_DONE.format(filename))

    def getAbility(self):

        ability = {'checkBox_b_title': False,
                   'checkBox_a_title': False,
                   'checkBox_target': False,
                   'checkBox_tail': False}

        if self.checkBox_b_title.isChecked():
            ability['checkBox_b_title'] = True

        if self.checkBox_a_title.isChecked():
            ability['checkBox_a_title'] = True

        if self.checkBox_target.isChecked():
            ability['checkBox_target'] = True

        if self.checkBox_tail.isChecked():
            ability['checkBox_tail'] = True

        return ability

    def dropEvent(self, event):
        '''
        松开鼠标
        '''
        if self.running:
            self.eventRestore("warning", STR_ERROR, STR_RUNNING_ERR)
            return

        file = event.mimeData().urls()[0].toLocalFile()

        if not is_word_file(file):
            self.eventRestore("warning", STR_ERROR, STR_DOC_TYPE_ERR)
            return
        self.running = True
        self.odf = ODFThread('\\'.join(file.split('/')),
                             self.conf, self.getAbility())
        self.odf.sig_err.connect(self.evenErrorCallback)
        self.odf.sig_label.connect(self.evenLabelCallback)
        self.odf.sig_label_info.connect(self.evenLabelInfoCallback)
        self.odf.sig_progress_init.connect(self.evenProgressInitCallback)
        self.odf.sig_progress_index.connect(self.evenProgressIndexCallback)
        self.odf.sig_done.connect(self.evenDoneCallback)
        self.odf.run()
        # odf.start()


class ODFThread(QThread):
    '''
    image to pdf 线程
    '''
    sig_err = pyqtSignal()
    sig_label = pyqtSignal(int, int)
    sig_label_info = pyqtSignal(str)
    sig_progress_init = pyqtSignal(int)
    sig_progress_index = pyqtSignal(int)
    sig_done = pyqtSignal(str)

    def __init__(self, file, conf, ability):
        '''
        构造函数
        '''

        super(ODFThread, self).__init__()
        self.file = file
        self.conf = conf
        self.ability = ability

    def getStepCount(self):
        return len(STEP)

    def run(self):
        '''
        线程启动函数
        '''
        d_to_dx = 0
        step_count = self.getStepCount()
        if get_file_extension(self.file) in [".doc", ".DOC"]:
            d_to_dx = 1
            step_count += 1

        try:
            if d_to_dx:
                self.sig_label.emit(1, step_count)
                self.sig_progress_init.emit(1)
                self.sig_progress_index.emit(1)
                self.sig_label_info.emit(STR_LABEL_DOCX)
                self.file = doc_to_docx(self.file)

            self.doc = Doc(self.file)
            self.doc.setSignal(self.sig_label, self.sig_label_info,
                               self.sig_progress_init, self.sig_progress_index, self.sig_done)
            self.doc.setDToDx(d_to_dx)
            self.doc.setStep(d_to_dx, step_count)
            self.doc.setAbility(self.ability)
            self.doc.setConf(self.conf['b_title']["line_count"], self.conf['a_title']
                             ["line_count"], self.conf['tail']["line_count"],)
            self.doc.run()
            self.doc.write()

        except Exception as e:
            print(e)
            self.sig_err.emit()


def load_conf(filename):
    with open(filename, 'r', encoding="utf-8") as f:
        return json.load(f)


def restore_conf(filename, data):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

    return data


def main():
    # '''
    # 主函数
    # '''
    app = QApplication(sys.argv)
    myWin = MyMainForm()
    myWin.show()
    myWin.setConfig(get_conf(myWin))

    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
